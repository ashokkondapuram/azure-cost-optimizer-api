#!/usr/bin/env python3
"""Generate comprehensive technical architecture document for CostOptimizeRecommender."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Shade a table cell with a background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_heading(paragraph_text, level=1):
    """Add a heading to the document."""
    p = doc.add_paragraph(paragraph_text, style=f'Heading {level}')
    p.runs[0].font.size = Pt([28, 24, 20][level-1])
    p.runs[0].font.bold = True
    return p

def add_text(text, size=11, bold=False, italic=False):
    """Add body text to the document."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    return p

def add_code_block(code_text):
    """Add a code/diagram block with monospace font."""
    p = doc.add_paragraph(code_text, style='Normal')
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(46, 117, 182)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add bottom border
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    """Add a table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        shade_cell(cell, '2E75B6')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_idx, row_data in enumerate(rows, 1):
        cells = table.rows[row_idx].cells
        for cell_idx, cell_text in enumerate(row_data):
            cells[cell_idx].text = cell_text

    # Set column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

def add_page_break():
    """Add a page break."""
    doc.add_page_break()

# Create document
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# ===== TITLE PAGE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("TECHNICAL ARCHITECTURE")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 78, 120)

subtitle = doc.add_paragraph("CostOptimizeRecommender Platform", style='Normal')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(24)
    run.font.bold = True

subtitle2 = doc.add_paragraph("Comprehensive System Design & Component Documentation", style='Normal')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle2.runs:
    run.font.size = Pt(14)
    run.font.italic = True

doc.add_paragraph()  # Spacing
doc.add_paragraph()

version = doc.add_paragraph("Version: 2.0", style='Normal')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER

date = doc.add_paragraph("Date: July 2026", style='Normal')
date.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle3 = doc.add_paragraph("Enterprise FinOps Platform with Azure Integration", style='Normal')
subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle3.runs:
    run.font.size = Pt(12)
    run.font.italic = True

add_page_break()

# ===== EXECUTIVE SUMMARY =====
create_heading("Executive Summary")
add_text("CostOptimizeRecommender is an enterprise FinOps platform designed to provide organizations with centralized Azure cost visibility, resource inventory management, and actionable optimization recommendations. The platform integrates seamlessly with Azure Cost Management APIs and Azure Resource Manager to deliver real-time cost insights grounded in actual billed costs.")

add_text("The system implements a modern three-tier architecture with:")
add_text("• React-based Single Page Application (SPA) frontend", 10)
add_text("• FastAPI REST API backend with enterprise features", 10)
add_text("• PostgreSQL relational database for persistent storage", 10)
add_text("• Optional Kubernetes telemetry integration for utilization insights", 10)
add_text("• JWT-based authentication with role-based access control (RBAC)", 10)
add_text("• Historical data persistence for trending and advanced analysis", 10)

add_page_break()

# ===== SYSTEM ARCHITECTURE OVERVIEW =====
create_heading("System Architecture Overview")
create_heading("1. Three-Tier Architecture Model", 2)

diagram1 = """
┌─────────────────────────────────────────────────────────────┐
│                   PRESENTATION LAYER                          │
│              React SPA Frontend (Node.js)                    │
│  Dashboard | Cost Explorer | Inventory | Recommendations   │
└──────────────────────────┬──────────────────────────────────┘
                           │ REST API / JSON
┌──────────────────────────▼──────────────────────────────────┐
│               APPLICATION LAYER                              │
│     FastAPI Backend (Python 3.13 + Uvicorn)                │
│  Cost Service | Resource Service | Optimization Engine     │
│  Authentication | Authorization | Business Logic           │
└──────────────────────────┬──────────────────────────────────┘
                           │ SQL Queries
┌──────────────────────────▼──────────────────────────────────┐
│                  DATA LAYER                                  │
│              PostgreSQL Database                            │
│  Cost Records | Resources | Findings | Configurations      │
└─────────────────────────────────────────────────────────────┘
"""
add_code_block(diagram1)

add_text("Each layer is independently scalable and can be deployed on separate infrastructure:")
add_text("• Frontend: Served as static files via CDN or App Service", 10)
add_text("• Backend: Stateless FastAPI servers enable horizontal scaling", 10)
add_text("• Database: PostgreSQL with connection pooling for concurrent access", 10)

doc.add_paragraph()

create_heading("2. Core System Components", 2)

diagram2 = """
┌────────────────────────────────────────────────────────────────┐
│                        EXTERNAL SYSTEMS                         │
│  Azure Cost Management | Azure Resource Manager | K8s Cluster  │
└──────────────────────────────┬─────────────────────────────────┘
                               │
┌──────────────────────────────▼─────────────────────────────────┐
│                      BACKEND API GATEWAY                        │
│  Authentication | Authorization | Request Routing | CORS       │
└──────────────┬──────────────────────┬──────────────────────────┘
               │                      │
┌──────────────▼──────────────┐  ┌───────▼─────────────────────┐
│   COST SERVICE              │  │  RESOURCE SERVICE           │
│ • Fetch costs from Azure    │  │ • Inventory all resources   │
│ • Normalize cost data       │  │ • Live ARM queries          │
│ • Store cost snapshots      │  │ • Cache management          │
│ • Cost analysis & trends    │  │ • Resource categorization   │
└──────────────┬──────────────┘  └───────┬─────────────────────┘
               │                          │
┌──────────────▼──────────────┐  ┌───────▼─────────────────────┐
│  OPTIMIZATION ENGINE        │  │  DATABASE SERVICE           │
│ • Rule-based analysis       │  │ • SQLAlchemy ORM            │
│ • Savings calculations      │  │ • Query optimization        │
│ • Finding generation        │  │ • Connection pooling        │
│ • Deduplication logic       │  │ • Migration management      │
└──────────────┬──────────────┘  └───────┬─────────────────────┘
               │                          │
               └──────────────┬───────────┘
                              │
                ┌─────────────▼────────────┐
                │   PostgreSQL Database    │
                │  (Persistent Storage)    │
                └──────────────────────────┘
"""
add_code_block(diagram2)

add_page_break()

# ===== DATA FLOW DIAGRAMS =====
create_heading("Data Flow Architecture")

create_heading("Cost Data Pipeline", 2)
cost_diagram = """
┌──────────────────────┐
│ Azure Cost           │
│ Management API       │
│ (v2024-08-01)        │
└──────────┬───────────┘
           │ Query costs for subscriptions
           │ (Timeframe, ResourceGroup filter)
┌──────────▼──────────────────┐
│ Cost Fetch Service           │
│ • Validate subscription      │
│ • Normalize cost data        │
│ • Aggregate costs            │
└──────────┬───────────────────┘
           │ Cost Record objects
┌──────────▼──────────────────┐
│ PostgreSQL Database          │
│ - CostRecord table           │
│ - Indexed by subscription    │
│ - Historical snapshots       │
└──────────┬───────────────────┘
           │ SQL query results
┌──────────▼──────────────────┐
│ REST API Endpoint            │
│ GET /api/costs/summary       │
│ GET /api/costs/by-service    │
└──────────┬───────────────────┘
           │ JSON response
┌──────────▼──────────────────┐
│ React Frontend               │
│ • Cost Explorer Charts       │
│ • Dashboard KPIs             │
│ • Trend Analysis             │
└──────────────────────────────┘
"""
add_code_block(cost_diagram)

create_heading("Resource Inventory Pipeline", 2)
resource_diagram = """
┌──────────────────────┐
│ Azure Resource       │
│ Manager API          │
│ (Live enumeration)   │
└──────────┬───────────┘
           │ List resources by type
           │ • VMs, Disks, AKS
           │ • Storage, DBs, Networking
┌──────────▼────────────────────┐
│ Resource Discovery Service     │
│ • Enumerate all resource types │
│ • Fetch resource properties    │
│ • Extract metadata             │
└──────────┬────────────────────┘
           │ Resource objects
┌──────────▼────────────────────┐
│ Resource Enrichment Module     │
│ • Categorize resources         │
│ • Link to cost data            │
│ • Calculate metrics            │
└──────────┬────────────────────┘
           │ Enriched resources
┌──────────▼────────────────────┐
│ PostgreSQL Database            │
│ - Resource Inventory Tables    │
│ - Index by type, subscription  │
│ - Cached snapshots             │
└──────────┬────────────────────┘
           │ Query results
┌──────────▼────────────────────┐
│ REST API Endpoints             │
│ GET /api/resources/{type}      │
│ GET /api/resources/{id}        │
└──────────┬────────────────────┘
           │ JSON response
┌──────────▼────────────────────┐
│ React Frontend                 │
│ • Resource Inventory Pages     │
│ • Filtered Search              │
│ • Resource Details             │
└────────────────────────────────┘
"""
add_code_block(resource_diagram)

create_heading("Optimization Analysis Pipeline", 2)
opt_diagram = """
┌─────────────────────────────────────┐
│ Admin: Trigger Optimization Run     │
│ POST /api/optimize/analyze          │
└──────────────┬──────────────────────┘
               │
┌──────────────▼──────────────────────┐
│ Optimization Engine                 │
│ • Load configuration rules           │
│ • Fetch cost data from DB            │
│ • Fetch resource inventory from DB   │
└──────────────┬──────────────────────┘
               │
┌──────────────▼──────────────────────────────┐
│ Analysis Execution                          │
│ For each resource type:                     │
│  • Apply specialized sub-engine             │
│  • Evaluate against rules                   │
│  • Calculate savings estimates              │
│  • Generate findings                        │
└──────────────┬──────────────────────────────┘
               │
┌──────────────▼──────────────────────────────┐
│ Finding Deduplication                       │
│ • Hash findings for uniqueness               │
│ • Merge similar findings                     │
│ • Track remediation status                   │
└──────────────┬──────────────────────────────┘
               │
┌──────────────▼──────────────────────────────┐
│ PostgreSQL Database                         │
│ - OptimizationFinding table                 │
│ - OptimizationRun table (history)          │
│ - Finding status & evidence                 │
└──────────────┬──────────────────────────────┘
               │
┌──────────────▼──────────────────────────────┐
│ REST API Endpoint                           │
│ GET /api/optimize/findings                  │
│ GET /api/optimize/findings/{id}             │
└──────────────┬──────────────────────────────┘
               │
┌──────────────▼──────────────────────────────┐
│ React Frontend                              │
│ • Recommendations Page                      │
│ • Savings Summary                           │
│ • Finding Details & Evidence                │
└─────────────────────────────────────────────┘
"""
add_code_block(opt_diagram)

add_page_break()

# ===== TECHNOLOGY STACK =====
create_heading("Technology Stack")

tech_data = [
    ["Frontend Framework", "React", "18.x", "Single Page Application (SPA)"],
    ["Build Tool", "Node.js / npm", "18.x", "Frontend build and package management"],
    ["Backend Framework", "FastAPI", "0.111+", "REST API with async support"],
    ["Server", "Uvicorn", "Latest", "ASGI application server"],
    ["Language", "Python", "3.13", "Backend development language"],
    ["Database", "PostgreSQL", "13+", "Relational data storage"],
    ["ORM", "SQLAlchemy", "2.0.36", "Database abstraction layer"],
    ["Migrations", "Alembic", "1.13.1", "Database schema versioning"],
    ["Authentication", "PyJWT", "2.8.0", "JWT token signing & validation"],
    ["Validation", "Pydantic", "2.10.6", "Data validation & serialization"],
    ["Azure SDK", "azure-identity", "Latest", "Managed Identity authentication"],
    ["Caching", "cachetools", "5.3.3", "LRU caching for performance"],
    ["Logging", "structlog", "24.1.0", "Structured logging framework"],
]

add_table(["Component", "Technology", "Version", "Purpose"], tech_data)

add_page_break()

# ===== DEPLOYMENT ARCHITECTURE =====
create_heading("Deployment Architecture")

create_heading("Azure App Service Deployment", 2)

app_service_diagram = """
┌──────────────────────────────────────────────────┐
│         Azure App Service (Linux)                │
│         Python 3.13 Runtime                      │
├──────────────────────────────────────────────────┤
│  Startup Configuration:                          │
│  • Oryx build system (auto-installs deps)        │
│  • Startup command:                              │
│    uvicorn app.main:app --host 0.0.0.0 --port 8000 │
│                                                   │
│  Health Check:                                   │
│  • Endpoint: /health/live                        │
│  • Interval: 60 seconds                          │
│  • Failure threshold: 3 attempts                 │
├──────────────────────────────────────────────────┤
│  Configuration (App Service settings):           │
│  • DATABASE_URL (PostgreSQL connection)          │
│  • AUTH_ENABLED (true in production)             │
│  • JWT_SECRET (secure random key)                │
│  • SETTINGS_ENCRYPTION_KEY (AES-256 key)        │
│  • CORS_ALLOWED_ORIGINS (frontend URLs)         │
│  • APP_ENV (prod/qa/dev)                         │
├──────────────────────────────────────────────────┤
│  Frontend Serving:                               │
│  • Built React SPA in /app/static/               │
│  • Served by FastAPI SPA route handler           │
│  • Cache control headers (static assets)         │
├──────────────────────────────────────────────────┤
│  Scaling:                                        │
│  • Horizontal: Multiple instances (stateless)    │
│  • Auto-scale based on CPU/memory metrics        │
│  • Database connection pooling                   │
└──────────────────────────────────────────────────┘
"""
add_code_block(app_service_diagram)

create_heading("CI/CD Pipeline (Azure Pipelines)", 2)

cicd_diagram = """
┌──────────────┐
│ Git Push     │
│ (main/feat)  │
└──────┬───────┘
       │
       ▼
┌──────────────────────────────┐
│ Azure Pipelines Trigger      │
│ (azure-pipelines.yml)        │
└──────┬───────────────────────┘
       │
       ▼
┌──────────────────────────────────────────┐
│ Build Stage:                             │
│ • Checkout code                          │
│ • Build frontend (React -> dist/)         │
│ • Run tests (backend/frontend)           │
│ • Create deployment package (zip)        │
│ • Publish artifacts                      │
└──────┬───────────────────────────────────┘
       │
       ▼
┌──────────────────────────────────────────┐
│ Deploy Stage:                            │
│ • Download artifacts                     │
│ • Configure App Service (settings)       │
│ • Run Alembic migrations (schema)        │
│ • Deploy zip to App Service              │
│ • Run smoke tests                        │
└──────┬───────────────────────────────────┘
       │
       ▼
┌──────────────────────────────────────────┐
│ Azure App Service Live                   │
│ (New version running)                    │
└──────────────────────────────────────────┘
"""
add_code_block(cicd_diagram)

add_page_break()

# ===== SECURITY ARCHITECTURE =====
create_heading("Security Architecture")

create_heading("Authentication & Authorization", 2)

add_text("JWT-based Authentication:")
add_text("• Login endpoint at /api/auth/login requires credentials", 10)
add_text("• JWT tokens issued with configurable expiration", 10)
add_text("• Rate limiting prevents brute force attacks", 10)
add_text("• Tokens validated on protected endpoints", 10)
add_text("• JWT_SECRET stored securely in production environments", 10)

add_text("")
add_text("Role-Based Access Control (RBAC):")

rbac_data = [
    ["Admin", "Read/Write all data, Execute analysis, Manage configs, Create users", "/api/optimize/analyze, /api/sync"],
    ["Viewer", "Read-only access to dashboards and reports", "/api/costs/*, /api/resources/*"],
]

add_table(["Role", "Permissions", "Protected Endpoints"], rbac_data)

create_heading("Data Protection", 2)

add_text("Sensitive data is protected through:")
add_text("• Settings Encryption: AES-256 encryption using SETTINGS_ENCRYPTION_KEY", 10)
add_text("• Database Security: PostgreSQL SSL/TLS connections, credentials in env vars", 10)
add_text("• K8s Agent Token: Shared secret token-based authentication", 10)
add_text("• JWT Secret: Cryptographically signed tokens with expiration", 10)
add_text("• CORS Control: Dynamic CORS middleware validates frontend origins", 10)

add_page_break()

# ===== API OVERVIEW =====
create_heading("API Endpoints Reference")

create_heading("Cost Management", 2)
cost_api_data = [
    ["GET /api/costs/summary", "JWT", "Cost summary for period"],
    ["GET /api/costs/by-service", "JWT", "Cost breakdown by Azure service"],
    ["GET /api/costs/daily", "JWT", "Daily cost trends"],
    ["GET /api/costs/by-resource-group", "JWT", "Cost by resource group"],
]
add_table(["Endpoint", "Auth", "Purpose"], cost_api_data)

create_heading("Resource Management", 2)
resource_api_data = [
    ["GET /api/resources", "JWT", "List resources with pagination"],
    ["GET /api/resources/{id}", "JWT", "Get details of specific resource"],
    ["GET /api/resources/{type}", "JWT", "List resources by type"],
    ["POST /api/sync", "JWT+Admin", "Trigger resource inventory sync"],
]
add_table(["Endpoint", "Auth", "Purpose"], resource_api_data)

create_heading("Optimization & Findings", 2)
opt_api_data = [
    ["POST /api/optimize/analyze", "JWT+Admin", "Execute optimization analysis run"],
    ["GET /api/optimize/findings", "JWT", "List all optimization findings"],
    ["GET /api/optimize/findings/{id}", "JWT", "Get finding details with evidence"],
    ["PUT /api/optimize/findings/{id}", "JWT+Admin", "Update finding status"],
]
add_table(["Endpoint", "Auth", "Purpose"], opt_api_data)

add_page_break()

# ===== MONITORING & PERFORMANCE =====
create_heading("Monitoring & Performance")

create_heading("Health Checks & Observability", 2)

add_text("Health Check Endpoints:")
add_text("• GET /health/live - Liveness probe (used by orchestrators)", 10)
add_text("• GET /api/status - Detailed system status", 10)

add_text("")
add_text("Structured Logging with structlog:")
add_text("• API request logging (timestamp, method, path, status, response time)", 10)
add_text("• Database query logging (table, rows affected, duration)", 10)
add_text("• External API calls (endpoint, status, duration)", 10)
add_text("• Error logging with stack traces and context", 10)
add_text("• Business logic events (findings count, savings, etc.)", 10)

create_heading("Performance Targets", 2)

perf_data = [
    ["API Response Time (p50)", "< 200ms", "List operations with caching"],
    ["API Response Time (p99)", "< 2000ms", "Worst-case scenarios"],
    ["Cost API Sync Duration", "< 30 seconds", "Update cost records"],
    ["Resource Sync Duration", "< 120 seconds", "ARM enumeration + DB upsert"],
    ["Optimization Run (per 5000 resources)", "< 5 minutes", "Full analysis"],
    ["Frontend Load Time (FCP)", "< 3 seconds", "First Contentful Paint"],
]
add_table(["Metric", "Target", "Notes"], perf_data)

add_page_break()

# ===== CONCLUSION =====
create_heading("Conclusion")

add_text("CostOptimizeRecommender is a production-ready enterprise platform built with modern technologies and best practices. The three-tier architecture ensures clear separation of concerns, independent scaling, and robust security.")

add_text("Key strengths of the architecture:")
add_text("• Scalability: Stateless backend enables horizontal scaling", 10)
add_text("• Security: JWT authentication, RBAC, Managed Identity, encrypted settings", 10)
add_text("• Reliability: Health checks, structured logging, database transactions", 10)
add_text("• Observability: Comprehensive logging, metrics, and status endpoints", 10)
add_text("• Flexibility: Docker containerization supports App Service, AKS, or on-premises", 10)
add_text("• Maintainability: Database migrations, clean code structure, API docs", 10)

add_text("")
add_text("The platform is designed to grow with organizational needs, from small pilots to enterprise-wide deployments managing thousands of resources and millions in cost optimization recommendations.")

add_text("")
doc_info_p = doc.add_paragraph()
doc_info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = doc_info_p.add_run("---")
run.font.size = Pt(10)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run("Document Version: 2.0 | Updated: July 2026")
run.font.size = Pt(10)
run.font.italic = True

footer2_p = doc.add_paragraph()
footer2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2_p.add_run("For questions or updates, refer to the project repository documentation.")
run.font.size = Pt(10)
run.font.italic = True

# Save document
output_path = "TECHNICAL_ARCHITECTURE.docx"
doc.save(output_path)
print(f"✓ Document created successfully: {output_path}")
